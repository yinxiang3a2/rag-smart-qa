"""
文档处理服务
负责从不同类型文件中提取文本内容
Word (.docx/.doc) -> python-docx
PDF -> PyMuPDF / pdfplumber
图片 -> OCR (预留)
"""
import logging
from pathlib import Path
from typing import List

import docx
import fitz  # PyMuPDF
import pdfplumber

from ..config import settings

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """
    文档处理器
    支持格式: PDF, Word(.docx/.doc), 纯文本, Markdown, 图片(OCR)
    """

    @staticmethod
    def extract_text(file_path: Path, file_type: str) -> str:
        """
        从文件中提取文本
        Args:
            file_path: 文件路径
            file_type: document / image
        Returns:
            提取的文本内容
        """
        ext = file_path.suffix.lower()

        if file_type == "document":
            if ext == ".pdf":
                return DocumentProcessor._extract_from_pdf(file_path)
            elif ext in [".docx", ".doc"]:
                return DocumentProcessor._extract_from_word(file_path)
            elif ext in [".txt", ".md"]:
                return DocumentProcessor._extract_from_text(file_path)
            else:
                raise ValueError(f"不支持的文档格式: {ext}")

        elif file_type == "image":
            return DocumentProcessor._extract_from_image(file_path)

        raise ValueError(f"未知的文件类型: {file_type}")

    @staticmethod
    def _extract_from_pdf(file_path: Path) -> str:
        """从PDF提取文本，优先 PyMuPDF，回退 pdfplumber"""
        logger.info(f"从PDF提取文本: {file_path.name}")

        # 方案1: PyMuPDF（优先，速度快）
        try:
            text = DocumentProcessor._extract_pdf_pymupdf(file_path)
            if text.strip():
                return text
        except Exception as e:
            logger.warning(f"PyMuPDF 提取失败: {e}，尝试 pdfplumber")

        # 方案2: pdfplumber（回退，对复杂布局支持更好）
        try:
            text = DocumentProcessor._extract_pdf_pdfplumber(file_path)
            if text.strip():
                return text
        except Exception as e:
            logger.warning(f"pdfplumber 提取失败: {e}")

        logger.error(f"PDF 文本提取失败: {file_path.name}")
        return ""

    @staticmethod
    def _extract_pdf_pymupdf(file_path: Path) -> str:
        """使用 PyMuPDF 从 PDF 提取文本"""
        doc = fitz.open(file_path)
        parts = []
        for page in doc:
            text = page.get_text()
            if text.strip():
                parts.append(text)
        doc.close()
        return "\n".join(parts)

    @staticmethod
    def _extract_pdf_pdfplumber(file_path: Path) -> str:
        """使用 pdfplumber 从 PDF 提取文本"""
        parts = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text and text.strip():
                    parts.append(text)
        return "\n".join(parts)

    @staticmethod
    def _extract_from_word(file_path: Path) -> str:
        """从 Word 文档提取文本（支持 .docx 和 .doc）"""
        logger.info(f"从Word提取文本: {file_path.name}")
        ext = file_path.suffix.lower()

        if ext == ".docx":
            # python-docx 直接支持 .docx
            doc = docx.Document(file_path)
            return "\n".join(para.text for para in doc.paragraphs if para.text.strip())

        elif ext == ".doc":
            # .doc 格式较旧，尝试用 python-docx 打开（可能失败）
            try:
                doc = docx.Document(file_path)
                return "\n".join(para.text for para in doc.paragraphs if para.text.strip())
            except Exception as e:
                logger.warning(f".doc 文件提取失败（格式较旧）: {e}，尝试返回原始内容")
                # .doc 是二进制格式，python-docx 不完全支持，返回提示
                return f"[注意] .doc 格式文本提取不完整，请另存为 .docx 后重新上传。原始文件名: {file_path.name}"

    @staticmethod
    def _extract_from_text(file_path: Path) -> str:
        """从纯文本/Markdown文件提取文本"""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()
        except UnicodeDecodeError:
            try:
                with open(file_path, "r", encoding="gbk") as f:
                    return f.read()
            except UnicodeDecodeError:
                with open(file_path, "r", encoding="latin-1") as f:
                    return f.read()

    @staticmethod
    def _extract_from_image(file_path: Path) -> str:
        """从图片提取文本(OCR) - 预留"""
        logger.info(f"OCR识别图片: {file_path.name}")
        return f"[OCR预留] 图片文件: {file_path.name}，OCR功能待集成"

    @staticmethod
    def chunk_text(text: str, chunk_size: int = None, overlap: int = None) -> List[str]:
        """
        将文本分块
        Args:
            text: 原始文本
            chunk_size: 分块大小，默认使用配置
            overlap: 重叠大小，默认使用配置
        Returns:
            文本块列表
        """
        if chunk_size is None:
            chunk_size = settings.CHUNK_SIZE
        if overlap is None:
            overlap = settings.CHUNK_OVERLAP

        if len(text) <= chunk_size:
            return [text]

        chunks = []
        start = 0
        while start < len(text):
            end = start + chunk_size
            chunk = text[start:end]
            chunks.append(chunk)
            start = end - overlap

        return chunks


# 全局单例
document_processor = DocumentProcessor()